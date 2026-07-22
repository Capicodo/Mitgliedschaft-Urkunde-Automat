import csv
import os
import random
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

PLACEHOLDER_PREFIX = "FIELD_"


def clean_path(raw: str) -> str:
    return os.path.expanduser(raw.strip().strip('"').strip("'"))


def resolve_path(raw: str) -> Path:
    path = Path(clean_path(raw))
    if not path.is_absolute():
        path = (Path.cwd() / path).resolve()
    return path.resolve()


def normalize_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def ask_for_path(prompt: str, must_exist: bool = True, is_file: bool = True) -> Path:
    while True:
        raw = input(prompt).strip()
        if not raw:
            print("Bitte einen Wert eingeben.")
            continue
        path = resolve_path(raw)
        if must_exist:
            if is_file and path.exists() and path.is_file():
                return path
            if not is_file and path.exists() and path.is_dir():
                return path
            print("Der angegebene Pfad wurde nicht gefunden. Bitte erneut eingeben.")
        else:
            return path


def resolve_column_name(row: dict, placeholder: str) -> str:
    token = placeholder[len(PLACEHOLDER_PREFIX):]
    normalized_token = normalize_text(token)

    if not row:
        return ""

    for header in row.keys():
        if normalize_text(header) == normalized_token:
            return header

    aliases = {
        "ID": ["mitgliedsnummer", "mitgliednummer", "id", "mitgliedid"],
        "NAME": ["mitglied", "name", "kundenname", "namekontakt"],
        "NSHARES": ["anteilgezeichnet", "anteilegezeichnet", "anzahlanteile", "shares", "nshares"],
        "TYPE": ["artdermitgliedschaft", "mitgliedschaft", "type", "mitgliedschaftsart"],
    }

    for alias in aliases.get(token.upper(), []):
        for header in row.keys():
            if normalize_text(header) == alias:
                return header

    return ""


def replace_placeholders_in_paragraph(paragraph, replacements: dict) -> None:
    for run in paragraph.runs:
        updated_text = run.text
        for placeholder, value in replacements.items():
            updated_text = updated_text.replace(placeholder, value)
        if updated_text != run.text:
            run.text = updated_text


def replace_placeholders_in_document(document: Document, replacements: dict) -> None:
    for paragraph in document.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, replacements)


def export_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    docx_path = docx_path.resolve()
    pdf_path = pdf_path.resolve()

    try:
        import win32com.client as win32

        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        doc.SaveAs2(str(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
        return
    except Exception as exc:
        print(f"Word-Export fehlgeschlagen: {exc}")

    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if libreoffice:
        try:
            subprocess.run(
                [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            converted_pdf = docx_path.with_suffix(".pdf")
            if converted_pdf.exists() and converted_pdf != pdf_path:
                converted_pdf.replace(pdf_path)
            return
        except Exception as exc:
            print(f"LibreOffice-Export fehlgeschlagen: {exc}")

    create_simple_pdf_from_docx(docx_path, pdf_path)


def create_simple_pdf_from_docx(docx_path: Path, pdf_path: Path) -> None:
    document = Document(docx_path)
    story = []
    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.fontName = "Helvetica"
    style.fontSize = 11

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            story.append(Paragraph(text, style))
            story.append(Spacer(1, 0.12 * inch))

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                story.append(Paragraph(" | ".join(cells), style))
                story.append(Spacer(1, 0.08 * inch))

    if not story:
        story.append(Paragraph("Keine Inhalte verfügbar.", style))

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    document_pdf = SimpleDocTemplate(str(pdf_path), pagesize=A4, leftMargin=0.75 * inch, rightMargin=0.75 * inch)
    document_pdf.build(story)


def main() -> None:
    print("=== Mitgliedschaft Urkunde Automat Prototype ===")
    print("Bitte geben Sie die Dateien ein. Sie können auch komplette Pfade mit Leerzeichen einfügen.")

    template_path = ask_for_path("Pfad zur DOCX-Vorlage: ", must_exist=True, is_file=True)
    csv_path = ask_for_path("Pfad zur CSV-Datei: ", must_exist=True, is_file=True)
    destination_folder = ask_for_path("Zielordner für die Urkunden: ", must_exist=False, is_file=False)

    destination_folder.mkdir(parents=True, exist_ok=True)

    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.DictReader(handle)
        rows = list(reader)

    if not rows:
        raise RuntimeError("Die CSV-Datei enthält keine Daten.")

    if not reader.fieldnames:
        raise RuntimeError("Die CSV-Datei enthält keine Spaltenüberschriften.")

    doc = Document(template_path)

    placeholders = []
    for paragraph in doc.paragraphs:
        for match in re.findall(r"FIELD_[A-Za-z0-9_]+", paragraph.text):
            placeholders.append(match)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for match in re.findall(r"FIELD_[A-Za-z0-9_]+", cell.text):
                    placeholders.append(match)

    placeholders = sorted(set(placeholders))
    if not placeholders:
        raise RuntimeError("In der Vorlage wurden keine Platzhalter gefunden (z. B. FIELD_NAME).")

    print(f"Gefundene Platzhalter: {', '.join(placeholders)}")

    selected_rows = random.sample(rows, k=min(50, len(rows)))

    for index, row in enumerate(selected_rows, start=1):
        
        template_copy = Document(template_path)
        replacements = {}
        for placeholder in placeholders:
            column_name = resolve_column_name(row, placeholder)
            value = row.get(column_name, "") if column_name else ""
            replacements[placeholder] = str(value or "")

        replace_placeholders_in_document(template_copy, replacements)

        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", f"{row.get('Mitglied', 'mitglied') or 'mitglied'}").strip("_")
        if not safe_name:
            safe_name = f"mitglied_{index}"

        docx_output = destination_folder / f"urkunde_{index:02d}_{safe_name}.docx"
        pdf_output = destination_folder / f"urkunde_{index:02d}_{safe_name}.pdf"

        template_copy.save(docx_output)
        export_to_pdf(docx_output, pdf_output)
        print(f"Erstellt: {docx_output.name} und {pdf_output.name}")

    print("Fertig. Die Dateien wurden im Zielordner gespeichert.")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nAbgebrochen.")
    except Exception as exc:
        print(f"Fehler: {exc}")
        sys.exit(1)
