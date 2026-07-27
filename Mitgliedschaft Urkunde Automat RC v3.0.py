import csv
import os
import random
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


import builtins
import datetime


class DualLogger:
    """Redirects stdout to both the console and a specified log file.

    Also handles logging user input.
    """

    def __init__(self, file_path: Path):
        self.terminal = sys.stdout
        self.file_path = file_path
        # Clear or create the log file at startup
        with open(self.file_path, "a", encoding="utf-8") as f:
            f.write(
                f"\n--- New Session Started at {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} ---\n\n"
            )

    def write(self, message: str):
        # Print to terminal
        self.terminal.write(message)
        # Append to log file
        with open(self.file_path, "a", encoding="utf-8") as f:
            f.write(message)

    def flush(self):
        # Needed for Python's sys.stdout interface compatibility
        self.terminal.flush()


def setup_logging(log_filename: str = "app_activity.log") -> Path:
    """Sets up automatic logging for all print statements and input prompts/responses."""
    log_path = Path.cwd() / log_filename

    # Redirect standard output to DualLogger
    logger = DualLogger(log_path)
    sys.stdout = logger

    # Wrap built-in input() to capture user inputs in the log file
    original_input = builtins.input

    def logged_input(prompt: str = "") -> str:
        user_response = original_input(prompt)
        # Log the prompt + user entry since stdout redirect doesn't automatically catch input() returns
        with open(log_path, "a", encoding="utf-8") as f:
            f.write(f"\n[USER INPUT] {user_response}\n")
        return user_response

    builtins.input = logged_input
    return log_path


PLACEHOLDER_PREFIX = "FIELD_"


def clean_path(raw: str) -> str:
    return os.path.expanduser(raw.strip().strip('"').strip("'"))


def resolve_path(raw: str) -> Path:
    path = Path(clean_path(raw))
    if not path.is_absolute():
        path = (Path.cwd() / path).resolve()
    return path.resolve()


def normalize_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def ask_for_path(prompt: str, must_exist: bool = True, is_file: bool = True) -> Path:
    while True:
        raw = input(prompt).strip()
        if not raw:
            print("Bitte einen Wert eingeben.")
            continue
        path = resolve_path(raw)
        if must_exist:
            if is_file and path.exists() and path.is_file():
                return path
            if not is_file and path.exists() and path.is_dir():
                return path
            print("Der angegebene Pfad wurde nicht gefunden. Bitte erneut eingeben.")
        else:
            return path


def ask_for_outlook_account() -> object:
    try:
        import win32com.client as win32
    except Exception as exc:
        raise RuntimeError(f"Outlook-Integration ist nicht verfügbar: {exc}")

    outlook = win32.Dispatch("Outlook.Application")
    accounts = list(outlook.Session.Accounts)
    if not accounts:
        raise RuntimeError("Es wurden keine Outlook-Konten gefunden.")

    print("Verfügbare Outlook-Konten:")
    for index, account in enumerate(accounts, start=1):
        display_name = getattr(account, "DisplayName", "") or ""
        smtp_address = getattr(account, "SmtpAddress", "") or ""
        print(f"{index}. {display_name} ({smtp_address})")

    while True:
        value = input("Bitte geben Sie den zu verwendenden Outlook-Account ein (Display-Name oder E-Mail): ").strip()
        if not value:
            print("Bitte einen Account angeben.")
            continue
        for account in accounts:
            display_name = getattr(account, "DisplayName", "") or ""
            smtp_address = getattr(account, "SmtpAddress", "") or ""
            user_name = getattr(account, "UserName", "") or ""
            if value.lower() in display_name.lower() or value.lower() in smtp_address.lower() or value.lower() in user_name.lower():
                return getattr(account, "SmtpAddress", "") or getattr(account, "DisplayName", "")
        print("Der angegebene Account wurde nicht gefunden. Bitte erneut eingeben.")


def read_rows(csv_path: Path):
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        sample = handle.read(4096)
        handle.seek(0)
        if ";" in sample and "," not in sample:
            reader = csv.DictReader(handle, delimiter=';')
        else:
            reader = csv.DictReader(handle)
        return list(reader), reader.fieldnames


def resolve_column_name(row: dict, placeholder: str) -> str:
    token = placeholder[len(PLACEHOLDER_PREFIX):]
    normalized_token = normalize_text(token)

    if not row:
        return ""

    for header in row.keys():
        if normalize_text(header) == normalized_token:
            return header

    aliases = {
        "ID": ["mitgliedsnummer"],
        "NAME": ["mitglied"],
        "NSHARES": ["anteileeingezahlt"],
        "TYPE": ["investierendesmitglied"],
        "DATE": ["mitgliedseit"]
    }

    for alias in aliases.get(token.upper(), []):
        normalized_alias = normalize_text(alias)
        for header in row.keys():
            if normalize_text(header) == normalized_alias:
                return header

    return ""

def convert_to_simple_date(date_val):
    """Converts a German date string (DD.MM.YYYY) into a format without leading zeros (D.M.YYYY).

    Example: '27.07.2026' -> '27.7.2026'
    """
    # Parse the string into a datetime object
    dt = datetime.datetime.strptime(date_val, "%d.%m.%Y")

    # Reformat using non-zero-padded values (%-d and %-m work on Unix systems,
    # but constructing it via attributes ensures cross-platform compatibility)
    return f"{dt.day}.{dt.month}.{dt.year}"

def resolve_placeholder_value(row: dict, placeholder: str) -> str:
    column_name = resolve_column_name(row, placeholder)
    if not column_name:
        return ""
    
    if placeholder.upper() == "FIELD_TYPE":
        raw_value = row.get(column_name, "")
        if normalize_text(str(raw_value)) in {"true"}:
            return "investierend"
        if normalize_text(str(raw_value)) in {"false"}:
            return "nutzend"
        raise Exception("Error: There was neither true nor false in the field for Investierendes Mitglied")
        
    if placeholder.upper() == "FIELD_DATE":
        raw_value = row.get(column_name, "")
        if not raw_value:
            return ""
        return convert_to_simple_date(raw_value)

    value = row.get(column_name, "")
    return str(value or "")


def get_template_for_row(row: dict, template_dir: Path) -> Path:
    value = str(row.get("Snglr-Plrl", "") or "").strip().upper()
    if value == "P":
        return template_dir / "Plur_Bestätigung-der-Mitgliedschaft-in-der-Piluweri-eG.docx"
    return template_dir / "Sing_Bestätigung-der-Mitgliedschaft-in-der-Piluweri-eG.docx"


def replace_placeholders_in_paragraph(paragraph, replacements: dict) -> None:
    for run in paragraph.runs:
        updated_text = run.text
        for placeholder, value in replacements.items():
            updated_text = updated_text.replace(placeholder, value)
        if updated_text != run.text:
            run.text = updated_text


def replace_placeholders_in_document(document: Document, replacements: dict) -> None:
    for paragraph in document.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, replacements)


def export_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    docx_path = docx_path.resolve()
    pdf_path = pdf_path.resolve()

    try:
        import win32com.client as win32

        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        doc.SaveAs2(str(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
        return
    except Exception as exc:
        print(f"Word-Export fehlgeschlagen: {exc}")

    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if libreoffice:
        try:
            subprocess.run(
                [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            converted_pdf = docx_path.with_suffix(".pdf")
            if converted_pdf.exists() and converted_pdf != pdf_path:
                converted_pdf.replace(pdf_path)
            return
        except Exception as exc:
            print(f"LibreOffice-Export fehlgeschlagen: {exc}")

    create_simple_pdf_from_docx(docx_path, pdf_path)


def create_simple_pdf_from_docx(docx_path: Path, pdf_path: Path) -> None:
    document = Document(docx_path)
    story = []
    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.fontName = "Helvetica"
    style.fontSize = 11

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            story.append(Paragraph(text, style))
            story.append(Spacer(1, 0.12 * inch))

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                story.append(Paragraph(" | ".join(cells), style))
                story.append(Spacer(1, 0.08 * inch))

    if not story:
        story.append(Paragraph("Keine Inhalte verfügbar.", style))

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    document_pdf = SimpleDocTemplate(str(pdf_path), pagesize=A4, leftMargin=0.75 * inch, rightMargin=0.75 * inch)
    document_pdf.build(story)


def send_pdf_via_outlook(
    account_identifier: str,  # Übergeben Sie hier die SMTP-Adresse oder den DisplayName
    pdf_path: Path, 
    recipient: str, 
    name: str, 
    is_plural: bool = False
) -> bool:
    try:
        import win32com.client as win32

        outlook = win32.Dispatch("Outlook.Application")
        
        # 1. Passendes Konto frisch aus der aktuellen Session holen
        target_account = None
        for acc in outlook.Session.Accounts:
            smtp_addr = getattr(acc, "SmtpAddress", "") or ""
            disp_name = getattr(acc, "DisplayName", "") or ""
            if account_identifier.lower() in smtp_addr.lower() or account_identifier.lower() in disp_name.lower():
                target_account = acc
                break

        if not target_account:
            raise RuntimeError(f"Konto '{account_identifier}' konnte in Outlook nicht aktiviert werden.")

        mail = outlook.CreateItem(0)

        # 2. Konto explizit & zuverlässig zuweisen
        mail.SendUsingAccount = target_account
        # Zuverlässiger Low-Level COM-Invoke für neuere Outlook/Office 2024 Versionen:
        mail._oleobj_.Invoke(*(64209, 0, 8, 0, target_account))

        mail.To = recipient
        mail.Subject = "Korrektur der Bestätigung der Mitgliedschaft"
        
        # In Outlook anzeigen, um Standard-Signatur zu laden
        mail.Display()

        if is_plural:
            body_text = (
                f"<p>Hallo {name},</p>"
                "<p>heute ist mal wieder der Wurm drin. Zum Glück nicht in unserem Gemüse und Obst, aber unsere IT funktioniert noch immer nicht gut genug. Bitte entschuldigt die (erneute) Panne und das falsche PDF, das Ihr bekommen habt. <br/>"
                "<p>Das PDF-Dokument im Anhang dieser Nachricht fasst die wesentlichen Informationen zu Eurer Mitgliedschaft in der Piluweri eG nun endlich korrekt zusammen. Manche Mitglieder haben sich das für Ihre Akten gewünscht. <br/>"
                "Ich nehme die fehlerhaften Nachrichten von heute Nachmittag auf meine Kappe und hoffe sehr, dass es jetzt passt. Trotzdem gilt: Bitte gebt uns Bescheid, wenn die Daten nicht stimmen oder wenn sich etwas Wichtiges daran ändert. </p>"
                "<p>Herzliche Grüße aus Eurer Gärtnerei<br/>"
                "Eure Piluweris</p>"
            )
        else:
            body_text = (
                f"<p>Hallo {name},</p>"
                "<p>heute ist mal wieder der Wurm drin. Zum Glück nicht in unserem Gemüse und Obst, aber unsere IT funktioniert noch immer nicht gut genug. Bitte entschuldige die (erneute) Panne und das falsche PDF, das Du bekommen hast. <br/>"
                "<p>Das PDF-Dokument im Anhang dieser Nachricht fasst die wesentlichen Informationen zu Deiner Mitgliedschaft in der Piluweri eG nun endlich korrekt zusammen. Manche Mitglieder haben sich das für Ihre Akten gewünscht. <br/>"
                "Ich nehme die fehlerhaften Nachrichten von heute Nachmittag auf meine Kappe und hoffe sehr, dass es jetzt passt. Trotzdem gilt: Bitte gib uns Bescheid, wenn die Daten nicht stimmen oder wenn sich etwas Wichtiges daran ändert. </p>"
                "<p>Herzliche Grüße aus Deiner Gärtnerei<br/>"
                "Michael Meuser</p>"
            )

        # Text vor die geladene Signatur einfügen
        mail.HTMLBody = body_text + "<br/><br/>" + mail.HTMLBody

        mail.Attachments.Add(str(pdf_path))
        mail.Send()
        return True

    except Exception as exc:
        print(f"E-Mail konnte nicht gesendet werden: {exc}")
        return False


def main() -> None:
    print("=== Mitgliedschaft Urkunde Automat RC v3.0 ===")
    print("Bitte geben Sie die Dateien ein. Sie können auch komplette Pfade mit Leerzeichen einfügen.")

    template_dir = ask_for_path("Ordner mit den DOCX-Vorlagen: ", must_exist=True, is_file=False)
    csv_path = ask_for_path("Pfad zur CSV-Datei: ", must_exist=True, is_file=True)
    destination_folder = ask_for_path("Zielordner für die Urkunden: ", must_exist=False, is_file=False)

    destination_folder.mkdir(parents=True, exist_ok=True)

    account = ask_for_outlook_account()

    rows, fieldnames = read_rows(csv_path)
    if not rows:
        raise RuntimeError("Die CSV-Datei enthält keine Daten.")
    if not fieldnames:
        raise RuntimeError("Die CSV-Datei enthält keine Spaltenüberschriften.")


    sent_count = 0
    failed_count = 0
    results = []

    
    for index, row in enumerate(rows, start=1):
        
        
        stadium_value = str(row.get("Stadium", "")).strip()
        if stadium_value != "Mitglied":
            print(f"\n[Übersprungen] Kein Mitglied (Stadium: '{stadium_value}') bei: {row.get('Mitglied', '')}")
            continue  # Correctly skips the rest of the current row!
        
        status_value = str(row.get("Status", "")).strip()
        if status_value != "OK":
            print(f"\n[Übersprungen] Mitglied nicht OK (Status: '{status_value}') bei: {row.get('Mitglied', '')}")
            continue  # Correctly skips the rest of the current row!    
            
        template_path = get_template_for_row(row, template_dir)
        if not template_path.exists():
            raise RuntimeError(f"Vorlage nicht gefunden: {template_path}")

        document = Document(template_path)

        placeholders = []
        for paragraph in document.paragraphs:
            for match in re.findall(r"FIELD_[A-Za-z0-9_]+", paragraph.text):
                placeholders.append(match)

        for table in document.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    for match in re.findall(r"FIELD_[A-Za-z0-9_]+", cell.text):
                        placeholders.append(match)

        placeholders = sorted(set(placeholders))
        if not placeholders:
            raise RuntimeError(f"In der Vorlage wurden keine Platzhalter gefunden: {template_path.name}")

        replacements = {}
        for placeholder in placeholders:
            value = resolve_placeholder_value(row, placeholder)
            replacements[placeholder] = str(value or "")

        replace_placeholders_in_document(document, replacements)

        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", f"{row.get('Mitglied', 'mitglied') or 'mitglied'}").strip("_")
        if not safe_name:
            safe_name = f"mitglied_{index}"

        docx_output = destination_folder / f"Bestätigung_der_Mitgliedschaft_{safe_name}.docx"
        pdf_output = destination_folder / f"Bestätigung_der_Mitgliedschaft_{safe_name}.pdf"

        document.save(docx_output)
        export_to_pdf(docx_output, pdf_output)

        recipient = row.get('E-Mail', '').strip()
        if not recipient:
            print(f"❌ [Fehler] Keine E-Mail-Adresse für {row.get('Mitglied', '')}. Überspringe Versand.")
            failed_count += 1
            results.append({
                "name": safe_name,
                "recipient": recipient,
                "sent": False,
                "pdf": pdf_output.name,
            })
            continue
        
        name = replacements.get("FIELD_NAME", row.get("Mitglied", "")) or ""
        is_plural=(row.get("Snglr-Plrl", "").strip().upper() == "P")
        success = send_pdf_via_outlook(account, pdf_output, recipient, name, is_plural)
        if success:
            sent_count += 1
        else:
            failed_count += 1

        results.append({
            "name": safe_name,
            "recipient": recipient,
            "sent": success,
            "pdf": pdf_output.name,
        })

        print(f"Erstellt und wird versendet: {pdf_output.name} -> {'erfolgreich ✅' if success else 'fehlgeschlagen'}")

    print("\n=== E-Mail Zusammenfassung ===")
    print(f"Gesendet: {sent_count}")
    print(f"Nicht gesendet: {failed_count}")
    print(f"Gesamt: {len(results)}")
    for result in results:
        status = "gesendet" if result["sent"] else "nicht gesendet"
        presymbol = "✅" if result["sent"] else "❌"
        print(f"{presymbol} {result['pdf']}: {status} an {result['recipient']}")


if __name__ == "__main__":
    log_file = setup_logging("session_log.txt")
    try:
        main()
        input("\nPress Enter to exit...")
    except KeyboardInterrupt:
        print("\nAbgebrochen.")
        input("\nPress Enter to exit...")
    except Exception as exc:
        print(f"Fehler: {exc}")
        input("\nPress Enter to exit...")
        sys.exit(1)
