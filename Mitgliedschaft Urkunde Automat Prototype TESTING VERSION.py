import csv
import os
import random
import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

PLACEHOLDER_PREFIX = "FIELD_"
DEFAULT_RECIPIENT = "dev@ite-pli.de"
MAX_CERTIFICATES = 100


def clean_path(raw: str) -> str:
    return os.path.expanduser(raw.strip().strip('"').strip("'"))


def resolve_path(raw: str) -> Path:
    path = Path(clean_path(raw))
    if not path.is_absolute():
        path = (Path.cwd() / path).resolve()
    return path.resolve()


def normalize_text(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", value.lower())


def ask_for_path(prompt: str, must_exist: bool = True, is_file: bool = True) -> Path:
    while True:
        raw = input(prompt).strip()
        if not raw:
            print("Bitte einen Wert eingeben.")
            continue
        path = resolve_path(raw)
        if must_exist:
            if is_file and path.exists() and path.is_file():
                return path
            if not is_file and path.exists() and path.is_dir():
                return path
            print("Der angegebene Pfad wurde nicht gefunden. Bitte erneut eingeben.")
        else:
            return path


def ask_for_outlook_account() -> object:
    try:
        import win32com.client as win32
    except Exception as exc:
        raise RuntimeError(f"Outlook-Integration ist nicht verfügbar: {exc}")

    outlook = win32.Dispatch("Outlook.Application")
    accounts = list(outlook.Session.Accounts)
    if not accounts:
        raise RuntimeError("Es wurden keine Outlook-Konten gefunden.")

    print("Verfügbare Outlook-Konten:")
    for index, account in enumerate(accounts, start=1):
        display_name = getattr(account, "DisplayName", "") or ""
        smtp_address = getattr(account, "SmtpAddress", "") or ""
        print(f"{index}. {display_name} ({smtp_address})")

    while True:
        value = input("Bitte geben Sie den zu verwendenden Outlook-Account ein (Display-Name oder E-Mail): ").strip()
        if not value:
            print("Bitte einen Account angeben.")
            continue
        for account in accounts:
            display_name = getattr(account, "DisplayName", "") or ""
            smtp_address = getattr(account, "SmtpAddress", "") or ""
            user_name = getattr(account, "UserName", "") or ""
            if value.lower() in display_name.lower() or value.lower() in smtp_address.lower() or value.lower() in user_name.lower():
                return account
        print("Der angegebene Account wurde nicht gefunden. Bitte erneut eingeben.")


def read_rows(csv_path: Path):
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        sample = handle.read(4096)
        handle.seek(0)
        if ";" in sample and "," not in sample:
            reader = csv.DictReader(handle, delimiter=';')
        else:
            reader = csv.DictReader(handle)
        return list(reader), reader.fieldnames


def resolve_column_name(row: dict, placeholder: str) -> str:
    token = placeholder[len(PLACEHOLDER_PREFIX):]
    normalized_token = normalize_text(token)

    if not row:
        return ""

    for header in row.keys():
        if normalize_text(header) == normalized_token:
            return header

    aliases = {
        "ID": ["mitgliedsnummer", "mitgliednummer", "id", "mitgliedid"],
        "NAME": ["mitglied", "name", "kundenname", "namekontakt"],
        "NSHARES": ["anteileeingezahlt", "anteilegezeichnet", "anteilgezeichnet", "anzahlanteile", "shares", "nshares"],
        "TYPE": ["investierendesmitglied"],
    }

    for alias in aliases.get(token.upper(), []):
        normalized_alias = normalize_text(alias)
        for header in row.keys():
            if normalize_text(header) == normalized_alias:
                return header

    return ""


def resolve_placeholder_value(row: dict, placeholder: str) -> str:
    column_name = resolve_column_name(row, placeholder)
    if not column_name:
        return ""

    if placeholder.upper() == "FIELD_TYPE":
        raw_value = row.get(column_name, "")
        if normalize_text(str(raw_value)) in {"wahr", "true", "1", "ja", "yes", "t", "y"}:
            return "investierend"
        return "nutzend"

    value = row.get(column_name, "")
    return str(value or "")


def get_template_for_row(row: dict, template_dir: Path) -> Path:
    value = str(row.get("Snglr-Plrl", "") or "").strip().upper()
    if value == "P":
        return template_dir / "Plrl_Bestätigung-der-Mitgliedschaft-in-der-Piluweri-eG.docx"
    return template_dir / "Snglr_Bestätigung-der-Mitgliedschaft-in-der-Piluweri-eG.docx"


def replace_placeholders_in_paragraph(paragraph, replacements: dict) -> None:
    for run in paragraph.runs:
        updated_text = run.text
        for placeholder, value in replacements.items():
            updated_text = updated_text.replace(placeholder, value)
        if updated_text != run.text:
            run.text = updated_text


def replace_placeholders_in_document(document: Document, replacements: dict) -> None:
    for paragraph in document.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, replacements)


def export_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    docx_path = docx_path.resolve()
    pdf_path = pdf_path.resolve()

    try:
        import win32com.client as win32

        word = win32.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(docx_path))
        doc.SaveAs2(str(pdf_path), FileFormat=17)
        doc.Close()
        word.Quit()
        return
    except Exception as exc:
        print(f"Word-Export fehlgeschlagen: {exc}")

    libreoffice = shutil.which("libreoffice") or shutil.which("soffice")
    if libreoffice:
        try:
            subprocess.run(
                [libreoffice, "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
                check=True,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            converted_pdf = docx_path.with_suffix(".pdf")
            if converted_pdf.exists() and converted_pdf != pdf_path:
                converted_pdf.replace(pdf_path)
            return
        except Exception as exc:
            print(f"LibreOffice-Export fehlgeschlagen: {exc}")

    create_simple_pdf_from_docx(docx_path, pdf_path)


def create_simple_pdf_from_docx(docx_path: Path, pdf_path: Path) -> None:
    document = Document(docx_path)
    story = []
    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.fontName = "Helvetica"
    style.fontSize = 11

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            story.append(Paragraph(text, style))
            story.append(Spacer(1, 0.12 * inch))

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                story.append(Paragraph(" | ".join(cells), style))
                story.append(Spacer(1, 0.08 * inch))

    if not story:
        story.append(Paragraph("Keine Inhalte verfügbar.", style))

    pdf_path.parent.mkdir(parents=True, exist_ok=True)
    document_pdf = SimpleDocTemplate(str(pdf_path), pagesize=A4, leftMargin=0.75 * inch, rightMargin=0.75 * inch)
    document_pdf.build(story)


def send_pdf_via_outlook(
    account: object, 
    pdf_path: Path, 
    recipient: str, 
    name: str, 
    is_plural: bool = False
) -> bool:
    try:
        import win32com.client as win32

        outlook = win32.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)
        mail.SendUsingAccount = account
        mail.To = recipient
        mail.Subject = "Bestätigung der Mitgliedschaft"
        
        # Displaying the item loads the user's default signature into mail.HTMLBody
        mail.Display()

        # Select message body based on singular ("Du") or plural ("Ihr") context
        if is_plural:
            body_text = (
                f"<p>Hallo {name},</p>"
                "<p>das PDF-Dokument im Anhang dieser Nachricht fasst die "
                "wesentlichen Informationen zu Eurer Mitgliedschaft in der Piluweri eG zusammen. "
                "So habt Ihr etwas für Eure Akten.<br/>"
                "Bitte gebt uns Bescheid, wenn die Daten nicht stimmen oder "
                "wenn sich etwas Wichtiges daran ändert.</p>"
                "<p>Herzliche Grüße aus Eurer Gärtnerei<br/>"
                "Eure Piluweris</p>"
            )
        else:
            body_text = (
                f"<p>Hallo {name},</p>"
                "<p>das PDF-Dokument im Anhang dieser Nachricht fasst die "
                "wesentlichen Informationen zu Deiner Mitgliedschaft in der Piluweri eG zusammen. "
                "So hast Du etwas für Deine Akten.<br/>"
                "Bitte gib uns Bescheid, wenn die Daten nicht stimmen oder "
                "wenn sich etwas Wichtiges daran ändert.</p>"
                "<p>Herzliche Grüße aus Deiner Gärtnerei<br/>"
                "Deine Piluweris</p>"
            )

        # Prepend custom body content before the default signature
        mail.HTMLBody = body_text + "<br/><br/>" + mail.HTMLBody

        mail.Attachments.Add(str(pdf_path))
        mail.Send()
        return True

    except Exception as exc:
        print(f"E-Mail konnte nicht gesendet werden: {exc}")
        return False


def main() -> None:
    print("=== Mitgliedschaft Urkunde Automat Pre-Alpha v2.0 ===")
    print("Bitte geben Sie die Dateien ein. Sie können auch komplette Pfade mit Leerzeichen einfügen.")

    template_dir = ask_for_path("Ordner mit den DOCX-Vorlagen: ", must_exist=True, is_file=False)
    csv_path = ask_for_path("Pfad zur CSV-Datei: ", must_exist=True, is_file=True)
    destination_folder = ask_for_path("Zielordner für die Urkunden: ", must_exist=False, is_file=False)

    destination_folder.mkdir(parents=True, exist_ok=True)

    account = ask_for_outlook_account()

    rows, fieldnames = read_rows(csv_path)
    if not rows:
        raise RuntimeError("Die CSV-Datei enthält keine Daten.")
    if not fieldnames:
        raise RuntimeError("Die CSV-Datei enthält keine Spaltenüberschriften.")

    selected_rows = random.sample(rows, k=min(MAX_CERTIFICATES, len(rows)))

    sent_count = 0
    failed_count = 0
    results = []

    singles_count = 0
    plurals_count = 0

    min_singles_required = 3
    min_plurals_required = 3
    
    for index, row in enumerate(selected_rows, start=1):
        
        if row.get("Snglr-Plrl") == "S":
            singles_count += 1
            if singles_count > min_singles_required:
                continue  # Skip processing if we have enough singles
        elif row.get("Snglr-Plrl") == "P":
            plurals_count += 1
            if plurals_count > min_plurals_required:
                continue  # Skip processing if we have enough plurals
            
        template_path = get_template_for_row(row, template_dir)
        if not template_path.exists():
            raise RuntimeError(f"Vorlage nicht gefunden: {template_path}")

        document = Document(template_path)

        placeholders = []
        for paragraph in document.paragraphs:
            for match in re.findall(r"FIELD_[A-Za-z0-9_]+", paragraph.text):
                placeholders.append(match)

        for table in document.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    for match in re.findall(r"FIELD_[A-Za-z0-9_]+", cell.text):
                        placeholders.append(match)

        placeholders = sorted(set(placeholders))
        if not placeholders:
            raise RuntimeError(f"In der Vorlage wurden keine Platzhalter gefunden: {template_path.name}")

        replacements = {}
        for placeholder in placeholders:
            value = resolve_placeholder_value(row, placeholder)
            replacements[placeholder] = str(value or "")

        replace_placeholders_in_document(document, replacements)

        safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", f"{row.get('Mitglied', 'mitglied') or 'mitglied'}").strip("_")
        if not safe_name:
            safe_name = f"mitglied_{index}"

        docx_output = destination_folder / f"urkunde_{index:02d}_{safe_name}.docx"
        pdf_output = destination_folder / f"urkunde_{index:02d}_{safe_name}.pdf"

        document.save(docx_output)
        export_to_pdf(docx_output, pdf_output)

        recipient = DEFAULT_RECIPIENT
        name = replacements.get("FIELD_NAME", row.get("Mitglied", "")) or ""
        is_plural=(row.get("Snglr-Plrl", "").strip().upper() == "P")
        success = send_pdf_via_outlook(account, pdf_output, recipient, name, is_plural)
        if success:
            sent_count += 1
        else:
            failed_count += 1

        results.append({
            "name": safe_name,
            "recipient": recipient,
            "sent": success,
            "pdf": pdf_output.name,
        })

        print(f"Erstellt und versendet: {pdf_output.name} -> {'erfolgreich' if success else 'fehlgeschlagen'}")

    print("\n=== E-Mail Zusammenfassung ===")
    print(f"Gesendet: {sent_count}")
    print(f"Nicht gesendet: {failed_count}")
    print(f"Gesamt: {len(results)}")
    for result in results:
        status = "gesendet" if result["sent"] else "nicht gesendet"
        print(f"- {result['pdf']}: {status} an {result['recipient']}")


if __name__ == "__main__":
    try:
        main()
        input("\nPress Enter to exit...")
    except KeyboardInterrupt:
        print("\nAbgebrochen.")
    except Exception as exc:
        print(f"Fehler: {exc}")
        sys.exit(1)
